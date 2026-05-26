from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("/Users/natalia/test/excel-analyzer/docs/案例书-办公小智快速落地案例.docx")


TITLE = "以跨团队“去信息差”机制提升协同质效"
SUBTITLE = "推动智能体应用快速落地\n——“大瓦特-办公小智”快速开发上线案例"


def set_east_asia_font(run, font_name):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_paragraph_format(paragraph, *, first_line=False, keep_next=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.25
    if first_line:
        fmt.first_line_indent = Cm(0.74)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.keep_together = False
    paragraph.paragraph_format.keep_with_next = keep_next


def add_para(doc, text="", *, style=None, bold=False, align=None, first_line=True, size=12, color=None, after=6):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.25
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    set_east_asia_font(r, "仿宋")
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_east_asia_font(r, "黑体")
    r.bold = True
    r.font.size = Pt(14 if level == 1 else 12.5)
    return p


def add_label_para(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.25
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(label)
    set_east_asia_font(r1, "黑体")
    r1.bold = True
    r1.font.size = Pt(12)
    r2 = p.add_run(text)
    set_east_asia_font(r2, "仿宋")
    r2.font.size = Pt(12)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, *, bold=False, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_east_asia_font(r, "仿宋")
    r.font.size = Pt(10.5)
    r.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shade_cell(cell, fill)


def set_table_width(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = OxmlElement("w:tcMar")
            for side in ("top", "left", "bottom", "right"):
                node = OxmlElement(f"w:{side}")
                node.set(qn("w:w"), "120")
                node.set(qn("w:type"), "dxa")
                tc_mar.append(node)
            tc_pr.append(tc_mar)


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.18)
section.right_margin = Cm(3.18)
section.header_distance = Cm(1.25)
section.footer_distance = Cm(1.25)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "仿宋"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
normal.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
r = p.add_run(TITLE)
set_east_asia_font(r, "黑体")
r.font.size = Pt(18)
r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
r = p.add_run(SUBTITLE)
set_east_asia_font(r, "黑体")
r.font.size = Pt(15)
r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run("（第X组）")
set_east_asia_font(r, "仿宋")
r.font.size = Pt(12)

add_para(doc, "小组成员：待补充（姓名  职务）", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, after=10)

add_heading(doc, "【基本信息】", level=1)
info = doc.add_table(rows=6, cols=4)
info.style = "Table Grid"
set_table_width(info)
rows = [
    ("案例类别", "管理类 / 重点工作类", "萃取属性", "任务达成类"),
    ("案例性质", "成功案例", "表现风格", "情景故事类 / 业务流程类"),
    ("成果名称", "大瓦特-办公小智", "项目时间", "2024年12月—2025年2月"),
    ("开发部门", "待补充", "本人角色", "团队主导与技术指导"),
    ("受众对象", "数字化项目负责人、产品经理、跨团队协同人员", "涉及团队", "3家公司、5个团队、6个智能体能力"),
    ("核心方法", "跨团队“去信息差”机制、短平快会议、问题清单、甘特图、阶段同步", "应用成效", "首月开放注册用户1200人，相关智能体全网推广使用"),
]
for row, values in zip(info.rows, rows):
    for idx, value in enumerate(values):
        set_cell_text(row.cells[idx], value, bold=idx in (0, 2), fill="F2F2F2" if idx in (0, 2) else None)

add_heading(doc, "【内容摘要】", level=1)
add_para(
    doc,
    "“大瓦特-办公小智”是人工智能技术赋能办公场景、促进智能体应用快速落地的重要实践。项目周期仅6周，且在推进中期面临平台架构重构、职责边界模糊、多团队信息割裂等压力。本案例围绕产品从方案重构到上线公测的全过程，总结通过短平快会议、问题清单、甘特图和阶段性向上同步等方式建立跨团队“去信息差”机制，推动3家公司、5个团队、6个智能体能力有序协同，最终实现按期交付并首月开放注册用户1200人的实践经验，为后续智能体产品开发、数字化项目落地和重点工作攻坚提供可复制参考。",
)
add_label_para(doc, "【关键词】：", "跨团队协同  去信息差  智能体应用  轻量化项目管理  快速交付")

add_heading(doc, "【案例目标】", level=1)
add_para(doc, "1. 提炼跨团队、多公司、多角色协同项目中的“去信息差”机制，帮助项目负责人在复杂任务中快速统一目标、边界、责任和风险。", first_line=False)
add_para(doc, "2. 总结短周期智能体产品开发中的轻量化项目管理方法，提升团队在沟通协调、节奏管控、问题闭环和资源整合方面的能力。", first_line=False)
add_para(doc, "3. 形成可复制、可推广的智能体应用快速落地经验，为后续人工智能办公产品建设、数字化项目推进和重点工作攻坚提供实践参考。", first_line=False)

add_heading(doc, "一、案例背景", level=1)
add_heading(doc, "（一）国家层面：人工智能加速赋能实体经济和组织治理", level=2)
add_para(
    doc,
    "党的二十大报告提出，要加快发展数字经济，促进数字经济和实体经济深度融合。近年来，国家持续推进数字中国建设，强调发挥数据要素和人工智能技术对产业升级、治理效能和生产方式变革的牵引作用。对电网企业而言，人工智能不再只是单点技术工具，而是支撑新型电力系统建设、提升企业经营管理质效和释放数据价值的重要抓手。办公场景智能体的建设，正是人工智能从技术探索走向日常业务应用的关键环节。",
)
add_heading(doc, "（二）公司层面：数字化转型进入智能化应用落地阶段", level=2)
add_para(
    doc,
    "南方电网公司持续推进数字电网建设和企业数字化转型，围绕业务在线化、数据资产化、能力平台化、服务智能化等方向，推动人工智能与电网生产、经营、管理、服务等场景深度融合。随着大模型和智能体技术快速发展，如何将已有数据、平台和业务能力转化为员工可直接使用的智能化工具，成为数字化部门和各业务单位共同面对的重要课题。",
)
add_heading(doc, "（三）项目层面：短周期、多团队、强协同成为交付关键", level=2)
add_para(
    doc,
    "2024年12月至2025年2月，借调至网公司数字化部期间，我作为“大瓦特-办公小智”专班项目主要负责人，承担复制深圳企智大脑开发智能体模式工作台、整合多方智能体能力、保障产品快速开发上线等任务。项目涉及数字运营公司、数字企业公司、人工智能公司等多方人员，以及网级数据中心、多平台运维、前端开发、数据接入、智能体能力建设等多个团队。项目周期仅6周，且在第3周出现平台架构需快速重构的情况，早期推进呈现职责边界模糊、信息传递不畅、开发节奏不一和需求反复调整等问题。",
)

add_heading(doc, "二、案例内容", level=1)
add_heading(doc, "（一）任务目标", level=2)
add_para(doc, "围绕“大瓦特-办公小智”6周内完成开发并上线公测的总体目标，项目提出以下配套子目标：")
add_para(doc, "1. 成果目标：复制深圳企智大脑开发智能体模式工作台，整合3家公司、5个团队、6个智能体能力，形成面向办公场景的智能体产品雏形。", first_line=False)
add_para(doc, "2. 应用目标：保障产品按期上线公测，面向全网用户开放注册和试用，推动相关智能体在办公场景中推广使用。", first_line=False)
add_para(doc, "3. 协同目标：建立跨团队“去信息差”机制，将目标、边界、责任、风险和节奏统一到同一张图、同一张表、同一套口径中。", first_line=False)
add_para(doc, "4. 可视化目标：通过问题清单、任务交付关键节点甘特图、日报周报和风险台账等工具，让项目进度、关键阻塞和待协调事项可视化、可追踪、可闭环。", first_line=False)

add_heading(doc, "（二）事件回顾", level=2)
add_para(
    doc,
    "项目启动初期，各团队均在各自职责范围内推进工作，但由于参与主体多、系统平台多、能力来源多，团队之间对“最终要做成什么样、谁负责什么、哪些能力必须上线、哪些问题需要上级拍板”的理解并不完全一致。项目推进到第3周时，领导认为当前系统平台架构与预期存在偏差，要求快速重构方案。此时项目周期已过半，如果仍按原有节奏推进，后续将面临更大的返工风险和上线压力。",
)
add_para(
    doc,
    "面对这一情况，我首先判断：项目真正的堵点并不只是技术开发，而是多方信息差没有被及时识别和消除。于是，我将工作重心从单纯催进度转向“去信息差、控节奏、提效率、稳方向”。一方面，通过短平快会议快速拉齐原型、架构、数据接入、智能体能力复用和上线范围；另一方面，通过问题清单和甘特图把责任人、协同方、完成时间和风险事项全部显性化，推动项目从“各自推进”转为“同屏协同”。",
)
add_para(
    doc,
    "在后续推进过程中，项目团队围绕进展、风险、待协调事项三类信息开展日报周报，不再做泛泛材料汇总；围绕上线节点倒排开发、联调、测试和公测准备，不再平均用力；围绕关键原型、架构和风险事项进行阶段性向上同步，不再等问题扩大后被动处理。最终，项目按期完成系统交付并上线公测，首月开放注册用户达到1200人，相关智能体在全网推广使用，部分智能体功能至今仍在使用。",
)

add_heading(doc, "（三）实施步骤", level=2)
add_heading(doc, "1. 技术支撑解决方案", level=2)
add_label_para(doc, "（1）存在问题：", "一是平台架构在项目中期被要求调整，原有方案需要快速推翻重构；二是6个智能体能力来自不同团队和系统，能力边界、接口方式、数据依赖和联调节奏不一致；三是前端开发、数据接入、智能体配置、平台运维等工作并行推进，任何一个环节滞后都可能影响整体上线。")
add_label_para(doc, "（2）原因分析：", "一是早期原型设计和平台架构未能在关键节点充分向上同步，导致领导预期和团队理解存在偏差；二是项目采取复制深圳企智大脑开发智能体模式工作台的方式推进，本地化过程中既要复用已有能力，又要适配网公司平台、数据和运维环境；三是各团队更熟悉自身模块，对跨团队依赖、交付边界和上线优先级缺少统一认知。")
add_label_para(doc, "（3）解决措施：", "一是快速梳理能力清单，明确哪些能力直接复用、哪些能力轻量改造、哪些需求暂缓优化，优先保障上线闭环；二是建立联调清单，将每个智能体对应的数据接口、前端页面、权限配置、运维支撑和责任人逐项拉通；三是用任务交付关键节点甘特图倒排方案确认、能力接入、页面联调、测试验证和上线准备，确保技术工作围绕关键路径推进；四是对低效环节及时调整协作方式和资源投入，避免局部问题拖慢整体交付。")

add_heading(doc, "2. 方向一致与协同权威问题", level=2)
add_label_para(doc, "（1）存在问题：", "一是各团队之间职责边界不够清晰，部分事项在团队间反复流转；二是关键信息停留在局部团队或个别沟通中，未形成统一事实底板；三是日报周报容易写成材料汇总，真正影响交付的风险和待协调事项没有被突出；四是项目探索性强，领导预期、业务体验和技术实现之间容易产生偏差。")
add_label_para(doc, "（2）原因分析：", "一是项目涉及3家公司、5个团队和多平台支撑，组织边界天然复杂；二是项目周期短，大家容易优先处理手头开发任务，而忽略跨团队信息对齐；三是缺少稳定的阶段性向上同步机制，导致关键原型、架构和风险事项没有在早期充分暴露。")
add_label_para(doc, "（3）解决措施：", "一是建立短平快会议机制，会议只围绕进展、风险、待协调事项展开，能点对点解决的问题不扩大讨论；二是建立“问题—责任人—协同方—截止时间—当前状态”五要素清单，避免“某团队负责”的模糊表达；三是统一日报周报口径，从“写得多”转向“解决得快”；四是在原型、架构、联调、上线前等关键节点主动向上同步，提前暴露需要拍板和协调的问题，争取决策支持。")

add_heading(doc, "3. 运营推广与上线保障问题", level=2)
add_label_para(doc, "（1）存在问题：", "一是产品需在6周内上线公测，方案重构后开发、联调、测试和公测准备时间进一步压缩；二是上线涉及权限开放、平台运维、数据支撑、用户注册和应用体验等多项工作，需要多方并行保障；三是智能体产品面向办公场景，用户体验和推广口径直接影响后续注册使用。")
add_label_para(doc, "（2）原因分析：", "一是公测不是内部演示，产品必须具备基本可用性和稳定性；二是多智能体能力成熟度不同，若不区分上线优先级，容易在非关键功能上消耗过多时间；三是项目团队前期主要精力集中在开发交付，对上线后的用户使用、推广节奏和持续优化也需要同步考虑。")
add_label_para(doc, "（3）解决措施：", "一是将上线任务拆分为必须完成、可复用、可延后优化三类，优先保障核心办公智能体能力可用；二是上线前集中梳理风险台账，对权限、运维、数据、页面、智能体能力等关键事项逐项确认；三是对可能影响上线的阻塞事项提前升级协调，避免问题在最后阶段集中出现；四是上线后跟踪注册用户和智能体使用情况，将项目经验沉淀为后续智能体产品开发和推广的可复用方法。")

add_heading(doc, "三、具体成效", level=1)
add_heading(doc, "（一）使用价值方面", level=2)
add_para(
    doc,
    "1. 办公用户侧：产品首月开放注册用户已达1200人，相关智能体在全网推广使用，部分智能体功能至今仍在使用。通过办公场景智能体，员工能够更便捷地获取智能化辅助能力，减少重复性信息检索和基础事务处理成本，推动人工智能从技术概念走向日常办公应用。",
    first_line=False,
)
add_para(
    doc,
    "2. 项目管理侧：项目从早期混乱和信息割裂转为有序推进，解决了多团队协作低效、问题反复转派和返工风险集中的问题。通过短平快会议、问题清单、甘特图和阶段同步，项目负责人能够更快识别关键责任人、关键风险和关键路径，提升复杂数字化项目的交付确定性。",
    first_line=False,
)
add_para(
    doc,
    "3. 组织能力侧：项目沉淀出跨团队“去信息差”机制和轻量化项目管理模型，可应用于后续人工智能应用开发、数字化项目落地和重点工作攻坚。该经验特别适用于周期紧、主体多、依赖强、变化快的复杂任务。",
    first_line=False,
)

add_heading(doc, "（二）品牌效益方面", level=2)
add_para(
    doc,
    "一是形成了人工智能赋能办公场景的内部示范案例。“大瓦特-办公小智”作为复制深圳企智大脑开发智能体模式工作台的重要成果，体现了公司在智能体应用快速落地方面的探索能力和组织协同能力。",
)
add_para(
    doc,
    "二是增强了智能体产品在全网推广使用的认知基础。产品上线后，相关智能体在全网范围内推广使用，部分功能持续保留，说明该类办公智能体具备一定应用黏性和场景价值。",
)
add_para(
    doc,
    "三是为后续宣传推广预留素材基础。当前素材暂未提供外部媒体报道数量、阅读量及具体报道链接，正式定稿时可补充网公司内网新闻、公众号推文、工作简报或专题汇报等宣传数据，进一步强化品牌效益和社会反响表述。",
)

add_heading(doc, "四、提问反思", level=1)
add_para(doc, "1. 当复杂项目推进到中期才暴露方向偏差时，项目负责人应优先催促开发，还是优先统一目标、边界和责任？", first_line=False)
add_para(doc, "2. 在多公司、多团队、多系统并行协作的项目中，哪些信息必须被显性化，才能避免反复返工和责任悬空？", first_line=False)
add_para(doc, "3. 面对短周期上线压力，如何在“快速交付”和“稳定可用”之间找到平衡点？", first_line=False)

add_heading(doc, "五、案例分析", level=1)
add_para(
    doc,
    "第一，复杂数字化项目的主要矛盾往往不是单一技术难题，而是目标、责任、节奏和风险之间的信息差。本案例中，项目前期出现混乱，并不是各团队不投入，而是各方看到的目标不完全一致、掌握的信息不完全一致、理解的优先级不完全一致。只有先把信息差显性化，后续开发和联调才具备稳定基础。",
)
add_para(
    doc,
    "第二，轻量化项目管理不是减少管理，而是减少无效管理。短平快会议、问题清单、日报周报和甘特图的价值，不在于增加流程，而在于让阻塞事项被看见、让责任人被明确、让风险能够提前处理。对于6周上线的项目，管理动作必须短、准、快，直接服务于问题解决和按期交付。",
)
add_para(
    doc,
    "第三，项目负责人既要懂技术，也要懂协同。在智能体产品开发中，负责人不能只关注功能实现，还要承担目标对齐者、信息枢纽者、节奏管控者、资源协调者和经验沉淀者等角色。尤其在方案重构、资源交叉和多线并行时，能否及时向上同步、横向协调和动态纠偏，直接决定项目能否从混乱走向有序。",
)

add_heading(doc, "六、总结升华", level=1)
add_para(
    doc,
    "本次案例是公司推动人工智能技术赋能办公场景、促进智能体应用快速落地的重要实践。通过该项目可以看到，数字化转型不仅需要技术能力，也需要把复杂任务组织起来、协同起来、闭环起来的管理能力。经验总结如下：",
)
add_para(
    doc,
    "（一）创新提出跨团队“去信息差”协同机制。针对多公司、多团队、多角色协同中目标不清、边界不明、责任不实等问题，通过短平快会议和问题清单，建立统一事实底板，实现从“各说各话”到“同屏协作”的转变。",
    first_line=False,
)
add_para(
    doc,
    "（二）系统性提出轻量化项目管理模型。围绕进展、风险、待协调事项三类关键信息，综合运用日报周报、甘特图、联调清单和风险台账等工具，形成适用于短周期数字化项目的节奏管控方法。",
    first_line=False,
)
add_para(
    doc,
    "（三）建立“阶段同步、风险前置、上线闭环”的快速交付模式。通过关键节点向上同步、核心能力优先复用、低效环节动态调整和上线前风险集中确认，保障智能体应用在短周期内完成开发、公测和推广，为后续重点工作攻坚提供可复制经验。",
    first_line=False,
)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run("案例书草稿｜大瓦特-办公小智")
set_east_asia_font(run, "仿宋")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(127, 127, 127)

doc.save(OUT)
print(OUT)
